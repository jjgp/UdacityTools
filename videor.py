import sys
import io
import re
import os
from docx import Document

# find *.mp4 | sed 's:\ :\\\ :g'| sed 's/^/file /' > fl.txt; ffmpeg -f concat -i fl.txt -c copy output.mp4; rm fl.txt

def addTranscriptToDocument(path, doc):
    f = open(path, 'r')
    o = io.StringIO()

    header = os.path.splitext(os.path.basename(path))[0]
    header += '\n\n'
    o.write(header.decode('utf-8'))

    section = re.compile('^[0-9]+$')
    timestamp = re.compile('^[0-9:,]+\\s-->\\s[0-9:,]+$')
    whitespace = re.compile('\\s+')
    for line in iter(f.readline, ''):
        if not section.match(line) and \
                not timestamp.match(line) and \
                not whitespace.match(line):
                    line = line.strip()
                    line += ' '
                    o.write(line.decode('utf-8'))

    o.write('\n\n'.decode('utf-8'))

    f.close()
    doc.add_paragraph(o.getvalue())
    o.close()

def ffmpegTheSubdir(path):
    for video in os.listdir(path):
        if video.endswith('.mp4'):
            addTranscriptToDocument(os.path.join(path, transcript), d)

def startWithTheRootDir(root):
    for subdir in os.listdir(root):
        if not subdir == ".DS_Store":
            ffmpegTheSubdir(os.path.join(root, subdir))

startWithTheRootDir(sys.argv[1])

