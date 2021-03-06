import sys
import io
import re
import os
from docx import Document


def addTranscriptToDocument(path, doc):
    f = open(path, 'r')
    o = io.StringIO()

    header = os.path.splitext(os.path.basename(path))[0]
    header += '\n\n'
    o.write(header.decode('utf-8'))

    section = re.compile('^[0-9]+(\\s?)+$')
    timestamp = re.compile('^[0-9:,]+\\s-->\\s[0-9:,]+(\\s?)+$')
    whitespace = re.compile('\\s+')
    for line in iter(f.readline, ''):
        if not section.match(line) and \
                not timestamp.match(line) and \
                not whitespace.match(line):
                    line = line.replace("&gt;", ">")
                    line = line.strip()
                    line += ' '
                    o.write(line.decode('utf-8'))

    o.write('\n\n'.decode('utf-8'))

    f.close()
    doc.add_paragraph(o.getvalue())
    o.close()


def docxifyTheSubdir(path):
    d = Document()
    number = lambda x: int(re.search('(^[0-9]+)', x).group(0))
    transcripts = os.listdir(path)
    transcripts = sorted(transcripts, key=number)

    for transcript in transcripts:
        if transcript.endswith('.srt'):
            addTranscriptToDocument(os.path.join(path, transcript), d)
    d.save(os.path.basename(path) + '.docx')


def startWithTheRootDir(root):
    for subdir in os.listdir(root):
        if not subdir == ".DS_Store":
            docxifyTheSubdir(os.path.join(root, subdir))


startWithTheRootDir(sys.argv[1])

